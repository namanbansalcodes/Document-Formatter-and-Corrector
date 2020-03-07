from tkinter import *
from tkinter import filedialog
import os
import docx
from itertools import islice
import pkg_resources
from symspellpy import SymSpell


class Engine():
    add = []  #stores selected file addresses
    filename = []  #stores selected file names
    doc = None  #doc object to work on current word file
    docf=None   #Final doc object
    n = 0  #number of lines in the doc file
    N = 0  #total number of files selected

    def __init__(self):
        self.loadsymspell()
    
    def inputfile(self, root):
        del self.add[::], self.filename[::]
        
        s = filedialog.askopenfilenames()

        if len(s) == 0:
            l21.config(text="0")
            l31.config(text="None")
        
        else:
            #seperating filename from the entire addresses
            for i in s:
                i = i[::-1]
                temp = 0
            
                for j in range(len(i)):
                    if i[j] == "/":
                        temp = j
                        break
            
                temp = len(i) - temp
                i = i[::-1]
            
                self.add.append(i[:temp])
                self.filename.append(i[temp:])
        
            self.N=len(self.add)

            l21.config(text=str(self.N))

            temp=""
            for i in self.filename:
                temp=temp+ i + "\n"
            temp=temp[:-1]

            l31.config(text=temp)
    
    def enginecontrol(self):
        for i in range(self.N):
            self.loaddoc(i)
            #symspellpy correction
            #punctuation correction
            #language-check correction

    def loadsymspell(self):
        sym_spell = SymSpell()
        dictionary_path = pkg_resources.resource_filename("symspellpy", "frequency_dictionary_en_82_765.txt")
        sym_spell.load_dictionary(dictionary_path, 0, 1)




    
    def loaddoc(self, i):
        os.chdir(self.add[i])  #changes the working directory
        
        self.doc = docx.Document(self.filename[i])
        self.doc.save("temp.docx")  #creating a copy of document so that original doc is not lost

        self.doc = docx.Document(self.filename[i])
        self.n = len(self.doc.paragraphs)
        
        self.docf=docx.Document()   #initializing result document 

e=Engine()
root = Tk()
root.title("Document Formatter and Corrector")
root.geometry("800x400")

#Left blue frame
f1 = Frame(root, width=200, height=400)
f1.config(background= "#b7e9fd")
f1.pack(side=LEFT)

#Right main frame
f2 = Frame(root, width=400, height=400)
f2.pack()

l0=Label(f2,text="")    #Blank label
l0.grid(row=0,column=0,columnspan=3)


#Browse button
b1 = Button(f2, text="Browse files",command= lambda:e.inputfile(root), width=60, height=1)
b1.grid(row=1, column=0, padx=2, pady=2, sticky=W, columnspan=3)
b1.config(relief="solid")

#Metadata1
l20=Label(f2,text="Number of files selected: ")
l20.grid(row=2,column=0,padx=2,pady=2,sticky=W)
l21=Label(f2,text="0")
l21.grid(row=2,column=1,padx=2,pady=2,sticky=W,columnspan=3)

#metadata2
l30=Label(f2,text="File Selected:")
l30.grid(row=3,column=0,padx=2,pady=2,sticky=W)
l31=Label(f2,text="None")
l31.grid(row=3, column=1, padx=2, pady=2,sticky=N+W)

#Go button
b41=Button(f2,text="Go!",command=lambda:e.enginecontrol(),width=60,height=1)
b41.grid(row=4, column=0, padx=2, pady=2, sticky=W, columnspan=3)
b41.config(relief="solid")

l50 = Label(f2, text="")
l50.grid(row=5, column=0, padx=2, pady=2, sticky=W)


root.mainloop()
